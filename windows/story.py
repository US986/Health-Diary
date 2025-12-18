"""
Модуль истории записей приложения "Дневник здоровья"
Содержит экран для просмотра, редактирования и экспорта медицинских записей
"""

# Стандартные библиотеки Python
import os
from datetime import datetime
import platform
import subprocess
import sys

from kivy.clock import Clock
# Библиотеки Kivy для создания GUI
from kivy.lang import Builder
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.screenmanager import Screen
from kivy.metrics import dp
from kivy.uix.scrollview import ScrollView

# Компоненты KivyMD (Material Design)
from kivymd.app import MDApp
from kivymd.uix.dialog import MDDialog
from kivymd.uix.textfield import MDTextField
from kivymd.uix.button import MDRaisedButton
from kivymd.uix.label import MDLabel
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.menu import MDDropdownMenu
from kivymd.uix.selectioncontrol import MDCheckbox
from kivymd.uix.list import TwoLineListItem

# Пользовательские модули
from database import get_connection, select_records_by_user, update_record, delete_record
from kv import REG_KV, PROFILE_KV, SETTINGS_KV, STORY_KV
from utils.rules import (
    validate_weight,
    validate_pressure_systolic,
    validate_pressure_diastolic,
    validate_pulse,
    validate_temperature,
    validate_notes
)

# Попытка импорта PIL для работы с изображениями
try:
    from PIL import Image as PILImage
    from PIL import ImageDraw

    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    PILImage = None
    ImageDraw = None

# Попытка импорта библиотек для экспорта документов
try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    import xlsxwriter

    XLSXWRITER_AVAILABLE = True
except ImportError:
    XLSXWRITER_AVAILABLE = False
    xlsxwriter = None

class StoryWindow(Screen):
    """
    Экран истории записей

    Позволяет пользователю:
    1. Просматривать историю медицинских записей
    2. Редактировать существующие записи
    3. Экспортировать записи в различные форматы (Word, Excel)
    4. Анализировать данные с помощью графиков
    """

    dialog = None  # Текущее диалоговое окно
    selected_records = {}  # Словарь выбранных записей для экспорта/удаления
    chart_menu = None  # Меню выбора типа графика
    selected_chart_type = "line"  # Выбранный тип графика по умолчанию
    search_query = ""  # Текст поиска
    all_records = []  # Все записи пользователя (для поиска)

    def __init__(self, **kwargs):
        """
        Инициализация экрана истории записей
        """
        super().__init__(**kwargs)
        # Запрашиваем разрешения на доступ к хранилищу на Android
        if platform == 'android':
            self.request_storage_permission()
        # Настраиваем меню выбора типа графика
        self.setup_chart_menu()

        # Расписание адаптации кнопок после загрузки UI
        Clock.schedule_once(self.adapt_ui_for_platform, 0.1)

    def adapt_ui_for_platform(self, dt):
        """
        Адаптирует UI под разные платформы
        """
        if platform == 'android':
            # На Android уменьшаем размер кнопок
            self.adapt_buttons_for_android()

    def adapt_buttons_for_android(self):
        """
        Адаптирует кнопки под Android
        """
        try:
            if hasattr(self, 'ids'):
                # Уменьшаем размер кнопок и расстояние между ними
                button_box = self.ids.get('button_box', None)
                if not button_box:
                    # Ищем контейнер с кнопками
                    for widget in self.walk():
                        if (hasattr(widget, 'children') and len(widget.children) > 0 and
                                isinstance(widget.children[0], BoxLayout) and
                                len(widget.children[0].children) > 0 and
                                hasattr(widget.children[0].children[0], 'icon')):
                            button_box = widget.children[0]
                            break

                if button_box:
                    # Уменьшаем высоту контейнера и отступы
                    button_box.height = dp(50)
                    button_box.padding = dp(5)
                    button_box.spacing = dp(8)

                    # Уменьшаем размер всех кнопок
                    for child in button_box.children:
                        if hasattr(child, 'size_hint_x'):
                            child.size_hint_x = None
                            child.width = dp(32)  # Уменьшенный размер
                            child.height = dp(32)
        except Exception as e:
            print(f"Ошибка адаптации кнопок: {e}")

    def setup_chart_menu(self):
        """
        Настраивает меню выбора типа графика для экспорта в Excel

        Создает список доступных типов графиков
        """
        chart_types = [
            {
                "text": "Линейный график",
                "viewclass": "OneLineListItem",
                "on_release": lambda x="line": self.set_chart_type(x)
            },
            {
                "text": "Столбчатая диаграмма",
                "viewclass": "OneLineListItem",
                "on_release": lambda x="bar": self.set_chart_type(x)
            },
            {
                "text": "Точечная диаграмма",
                "viewclass": "OneLineListItem",
                "on_release": lambda x="scatter": self.set_chart_type(x)
            },
            {
                "text": "Комбинированный график",
                "viewclass": "OneLineListItem",
                "on_release": lambda x="combo": self.set_chart_type(x)
            },
        ]

        # Создаем меню для выбора типа графика
        self.chart_menu = MDDropdownMenu(
            caller=None,  # Caller будет установлен при открытии
            items=chart_types,
            width_mult=4,  # Ширина меню (4 * стандартная ширина)
        )

    def set_chart_type(self, chart_type):
        """
        Устанавливает выбранный тип графика

        Args:
            chart_type (str): Тип графика ('line', 'bar', 'scatter', 'combo')
        """
        self.selected_chart_type = chart_type  # Сохраняем выбор
        self.chart_menu.dismiss()  # Закрываем меню
        # Обновляем текст кнопки, если она существует
        if hasattr(self, 'chart_type_button'):
            self.chart_type_button.text = f"Тип графика: {self.get_chart_type_name(chart_type)}"

    def get_chart_type_name(self, chart_type):
        """
        Возвращает читаемое имя типа графика

        Args:
            chart_type (str): Код типа графика

        Returns:
            str: Читаемое имя типа графика
        """
        names = {
            "line": "Линейный",
            "bar": "Столбчатый",
            "scatter": "Точечный",
            "combo": "Комбинированный"
        }
        return names.get(chart_type, "Неизвестно")

    def open_chart_menu(self, button):
        """
        Открывает меню выбора типа графика

        Args:
            button: Кнопка, которая вызвала меню
        """
        self.chart_menu.caller = button  # Устанавливаем кнопку-вызыватель
        self.chart_menu.open()  # Открываем меню

    def request_storage_permission(self):
        """
        Запрашивает разрешения на доступ к хранилищу на Android

        Необходимо для экспорта файлов в системные директории
        """
        try:
            if platform == 'android':
                from android.permissions import request_permissions, Permission
                # Запрашиваем разрешения на чтение и запись во внешнее хранилище
                request_permissions([Permission.WRITE_EXTERNAL_STORAGE, Permission.READ_EXTERNAL_STORAGE])
        except Exception as e:
            print(f"Ошибка при запросе разрешений: {e}")

    def get_export_directory(self):
        """
        Возвращает директорию для экспорта файлов

        Returns:
            str: Путь к директории для экспорта
        """
        try:
            if platform == 'android':
                # На Android используем папку Downloads
                from jnius import autoclass
                Environment = autoclass('android.os.Environment')
                downloads_dir = Environment.getExternalStoragePublicDirectory(Environment.DIRECTORY_DOWNLOADS)
                export_dir = os.path.join(downloads_dir.getAbsolutePath(), "HealthDiary")
            elif sys.platform == 'win32':  # Windows
                # На Windows используем папку Documents
                import winreg
                try:
                    # Получаем путь к папке Documents
                    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                                         r"Software\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders")
                    docs_path = winreg.QueryValueEx(key, "Personal")[0]
                    winreg.CloseKey(key)
                    export_dir = os.path.join(docs_path, "HealthDiary")
                except:
                    # Резервный вариант: папка в домашней директории
                    export_dir = os.path.join(os.path.expanduser("~"), "Documents", "HealthDiary")
            elif sys.platform == 'darwin':  # macOS
                export_dir = os.path.join(os.path.expanduser("~"), "Documents", "HealthDiary")
            else:  # Linux и другие
                export_dir = os.path.join(os.path.expanduser("~"), "HealthDiary")

            # Создаем директорию, если она не существует
            os.makedirs(export_dir, exist_ok=True)
            return export_dir

        except Exception as e:
            print(f"Ошибка при получении директории экспорта: {e}")
            # Резервный вариант: используем директорию данных приложения
            app_dir = os.path.join(MDApp.get_running_app().user_data_dir, "HealthDiaryExports")
            os.makedirs(app_dir, exist_ok=True)
            return app_dir

    def on_pre_enter(self):
        """
        Метод, вызываемый перед переходом на экран истории

        Загружает историю записей пользователя
        """
        self.load_story()

    def load_story(self, search_query=None):
        """
        Загружает историю записей пользователя из базы данных

        Отображает записи в списке с возможностью выбора для экспорта

        Args:
            search_query (str, optional): Текст для поиска записей
        """
        # Получаем ID текущего пользователя
        user_id = MDApp.get_running_app().get_user_id()
        if not user_id:
            return

        try:
            # Подключаемся к базе данных
            conn = get_connection()

            cursor = conn.cursor()
            # Загружаем записи пользователя, отсортированные по дате (новые сверху)
            records = select_records_by_user(conn, user_id)

            # Сохраняем все записи для поиска
            self.all_records = list(records) if records else []

            # Фильтруем записи, если есть поисковый запрос
            if search_query and search_query.strip():
                filtered_records = self.filter_records(self.all_records, search_query.strip())
                records = filtered_records
            else:
                records = self.all_records

            # Получаем контейнер для списка записей
            story_list = self.ids.container
            story_list.clear_widgets()  # Очищаем предыдущие записи
            self.selected_records = {}  # Сбрасываем выбранные записи

            if records:
                # Если есть записи - отображаем их
                for record in records:
                    record_id = record[0]  # ID записи
                    # Форматируем дату для отображения
                    record_date = self.format_display_date(record[7])

                    # Формируем тексты для отображения
                    primary_text = f"Дата: {record_date}"
                    secondary_text = (f"Вес: {record[1] if record[1] else 'Н/Д'} кг, "
                                      f"Давление: {record[2] if record[2] else 'Н/Д'}/{record[3] if record[3] else 'Н/Д'}, "
                                      f"Пульс: {record[4] if record[4] else 'Н/Д'}, "
                                      f"Темп.: {record[5] if record[5] else 'Н/Д'}°C")

                    # Создаем контейнер для записи (чекбокс + текст)
                    record_container = MDBoxLayout(
                        orientation='horizontal',
                        adaptive_height=True,
                        spacing=dp(10),
                        padding=dp(5)
                    )

                    # Создаем чекбокс для выбора записи
                    checkbox = MDCheckbox(
                        size_hint=(None, None),
                        size=(dp(40), dp(40)),
                        active=False  # По умолчанию не выбран
                    )

                    # Создаем элемент списка с текстом записи
                    list_item = TwoLineListItem(
                        text=primary_text,
                        secondary_text=secondary_text
                    )

                    # Привязываем обработчик клика для редактирования
                    list_item.bind(on_release=lambda x, rec=record: self.open_edit_form(rec))

                    # Привязываем обработчик изменения состояния чекбокса
                    checkbox.bind(
                        active=lambda instance, value, rec_id=record_id: self.on_checkbox_active(instance, value,
                                                                                                 rec_id))

                    # Добавляем элементы в контейнер
                    record_container.add_widget(checkbox)
                    record_container.add_widget(list_item)
                    story_list.add_widget(record_container)

                    # Сохраняем данные записи в словарь
                    self.selected_records[record_id] = {
                        'container': record_container,
                        'checkbox': checkbox,
                        'list_item': list_item,
                        'record': record,
                        'selected': False  # Флаг выбора
                    }
            else:
                # Если записей нет - показываем сообщение
                if search_query and search_query.strip():
                    empty_label = MDLabel(
                        text=f"Записей по запросу '{search_query}' не найдено",
                        halign="center",
                        theme_text_color="Secondary",
                        font_style="H6",
                        size_hint_y=None,
                        height=dp(60)
                    )
                else:
                    empty_label = MDLabel(
                        text="История записей пуста",
                        halign="center",
                        theme_text_color="Secondary",
                        font_style="H6",
                        size_hint_y=None,
                        height=dp(60)
                    )
                story_list.add_widget(empty_label)

        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при загрузке истории: {str(e)}")
        finally:
            if 'conn' in locals() and conn:
                conn.close()

    def filter_records(self, records, search_query):
        """
        Фильтрует записи по поисковому запросу

        Args:
            records: Список записей
            search_query: Текст для поиска

        Returns:
            list: Отфильтрованные записи
        """
        filtered = []
        search_query_lower = search_query.lower()

        for record in records:
            record_id = record[0]
            weight = str(record[1]) if record[1] else ""
            pressure_sys = str(record[2]) if record[2] else ""
            pressure_dia = str(record[3]) if record[3] else ""
            pulse = str(record[4]) if record[4] else ""
            temperature = str(record[5]) if record[5] else ""
            notes = str(record[6]) if record[6] else ""
            record_date = self.format_display_date(record[7])

            # Проверяем наличие поискового запроса в различных полях
            if (search_query_lower in weight.lower() or
                    search_query_lower in pressure_sys.lower() or
                    search_query_lower in pressure_dia.lower() or
                    search_query_lower in pulse.lower() or
                    search_query_lower in temperature.lower() or
                    search_query_lower in notes.lower() or
                    search_query_lower in record_date.lower()):
                filtered.append(record)

        return filtered

    def on_search(self, instance, value):
        """
        Обработчик изменения текста в поле поиска

        Args:
            instance: Поле ввода
            value: Текущее значение поля
        """
        self.search_query = value
        self.load_story(search_query=value)

    def format_display_date(self, date_value):
        """
        Форматирует дату в соответствии с настройками пользователя

        Args:
            date_value: Значение даты (может быть строкой или datetime)

        Returns:
            str: Отформатированная дата
        """
        try:
            app = MDApp.get_running_app()
            date_format = app.user_settings.get('date_format', 'dd-mm-yyyy')

            # Преобразуем строку в datetime, если нужно
            if isinstance(date_value, str):
                try:
                    # Пробуем разные форматы дат
                    for fmt in ["%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d-%m-%Y"]:
                        try:
                            date_value = datetime.strptime(date_value, fmt)
                            break
                        except ValueError:
                            continue
                except:
                    return date_value  # Если не удалось преобразовать, возвращаем как есть

            # Форматируем дату в соответствии с настройками
            if date_format == 'dd-mm-yyyy':
                return date_value.strftime("%d-%m-%Y")
            elif date_format == 'mm-dd-yyyy':
                return date_value.strftime("%m-%d-%Y")
            elif date_format == 'yyyy-mm-dd':
                return date_value.strftime("%Y-%m-%d")
            else:
                return date_value.strftime("%d-%m-%Y")  # Формат по умолчанию

        except Exception as e:
            print(f"Ошибка форматирования даты: {e}")
            if isinstance(date_value, str):
                return date_value
            else:
                return date_value.strftime("%d-%m-%Y")  # Формат по умолчанию

    def on_checkbox_active(self, checkbox, value, record_id):
        """
        Обработчик изменения состояния чекбокса

        Args:
            checkbox: Объект чекбокса
            value (bool): Новое состояние чекбокса
            record_id (int): ID записи
        """
        if record_id in self.selected_records:
            record_data = self.selected_records[record_id]
            record_data['selected'] = value  # Обновляем флаг выбора
            # Подсвечиваем запись, если она выбрана
            record_data['list_item'].md_bg_color = (0.9, 0.9, 1, 0.3) if value else (1, 1, 1, 1)

    def get_selected_records(self):
        """
        Возвращает список выбранных записей

        Returns:
            list: Список выбранных записей
        """
        return [record_data['record'] for record_data in self.selected_records.values() if record_data['selected']]

    def get_selected_record_ids(self):
        """
        Возвращает список ID выбранных записей

        Returns:
            list: Список ID выбранных записей
        """
        return [record_data['record'][0] for record_data in self.selected_records.values() if record_data['selected']]

    def add_new_record(self):
        """
        Переход к экрану добавления новой записи
        """
        self.manager.current = "options"

    def open_edit_form(self, record):
        """
        Открывает форму редактирования записи

        Args:
            record: Данные записи для редактирования
        """
        self.edit_record_dialog(record)

    def edit_record_dialog(self, record):
        """
        Создает диалоговое окно для редактирования записи

        Args:
            record: Данные записи для редактирования
        """
        # Создаем контейнер для полей ввода
        dialog_box = MDBoxLayout(
            orientation="vertical",
            spacing=dp(10),
            padding=dp(10),
            adaptive_height=True
        )

        # Определяем поля для редактирования
        fields = [
            ("Вес", str(record[1]) if record[1] is not None else ""),
            ("Систолическое давление", str(record[2]) if record[2] is not None else ""),
            ("Диастолическое давление", str(record[3]) if record[3] is not None else ""),
            ("Пульс", str(record[4]) if record[4] is not None else ""),
            ("Температура", str(record[5]) if record[5] is not None else ""),
            ("Заметки", str(record[6]) if record[6] is not None else "")
        ]

        # Создаем поля ввода
        inputs = []
        for hint, text in fields:
            input_field = MDTextField(hint_text=hint, text=text, mode="rectangle")
            if hint == "Заметки":
                input_field.multiline = True  # Поле заметок многострочное
            dialog_box.add_widget(input_field)
            inputs.append(input_field)

        # Создаем диалоговое окно
        self.dialog = MDDialog(
            title="Редактировать запись",
            type="custom",
            content_cls=dialog_box,
            buttons=[
                MDRaisedButton(
                    text="Отмена",
                    md_bg_color=(0.7, 0.7, 0.7, 1),
                    on_release=lambda _: self.dialog.dismiss()
                ),
                MDRaisedButton(
                    text="Сохранить",
                    md_bg_color=(0, 0, 1, 1),
                    on_release=lambda _: self.save_edit_record(record[0], *[inp.text for inp in inputs])
                ),
            ],
        )
        self.dialog.open()

    def save_edit_record(self, record_id, weight, pressure_systolic, pressure_diastolic, pulse, temperature, notes):
        """
        Сохраняет отредактированную запись в базу данных

        Args:
            record_id (int): ID записи
            weight (str): Вес
            pressure_systolic (str): Систолическое давление
            pressure_diastolic (str): Диастолическое давление
            pulse (str): Пульс
            temperature (str): Температура
            notes (str): Заметки
        """
        try:
            # Валидируем все поля
            weight = validate_weight(weight)
            pressure_systolic = validate_pressure_systolic(pressure_systolic)
            pressure_diastolic = validate_pressure_diastolic(pressure_diastolic)
            pulse = validate_pulse(pulse)
            temperature = validate_temperature(temperature)
            notes = validate_notes(notes)

            # Сохраняем изменения в базе данных
            conn = get_connection()

            update_record(conn, record_id, weight, pressure_systolic, pressure_diastolic, pulse, temperature, notes)

            # Закрываем диалог и обновляем список
            self.dialog.dismiss()
            self.load_story()
            self.show_message("Успех", "Запись успешно обновлена")

        except ValueError as ve:
            # Ошибка валидации данных
            self.show_message("Ошибка ввода", str(ve))
        except Exception as e:
            # Общая ошибка
            self.show_message("Ошибка", f"Ошибка при сохранении данных: {str(e)}")
        finally:
            if 'conn' in locals() and conn:
                conn.close()

    def show_message(self, title, text):
        """
        Показывает диалоговое окно с сообщением

        Args:
            title (str): Заголовок сообщения
            text (str): Текст сообщения
        """
        dialog = MDDialog(
            title=title,
            text=text,
            buttons=[
                MDRaisedButton(
                    text="OK",
                    md_bg_color=(0, 0, 1, 1),
                    on_release=lambda _: dialog.dismiss()
                )
            ]
        )
        dialog.open()

    def on_arrow_pressed(self):
        """
        Переход к экрану профиля пользователя
        """
        self.manager.current = "profile"

    def safe_convert_to_float(self, value):
        """
        Безопасно преобразует значение в float

        Args:
            value: Значение для преобразования

        Returns:
            float или None: Преобразованное значение или None в случае ошибки
        """
        try:
            return float(value) if value is not None else None
        except (TypeError, ValueError):
            return None

    def safe_convert_to_int(self, value):
        """
        Безопасно преобразует значение в int

        Args:
            value: Значение для преобразования

        Returns:
            int или None: Преобразованное значение или None в случае ошибки
        """
        try:
            return int(float(value)) if value is not None else None
        except (TypeError, ValueError):
            return None

    def delete_selected_records(self):
        """
        Удаляет выбранные записи из базы данных
        """
        selected_ids = self.get_selected_record_ids()

        if not selected_ids:
            self.show_message("Внимание", "Сначала выберите записи для удаления")
            return

        # Создаем диалог подтверждения удаления
        dialog = MDDialog(
            title="Подтверждение удаления",
            text=f"Вы уверены, что хотите удалить {len(selected_ids)} выбранных записей?\nЭто действие нельзя отменить.",
            buttons=[
                MDRaisedButton(
                    text="Отмена",
                    md_bg_color=(0.7, 0.7, 0.7, 1),
                    on_release=lambda _: dialog.dismiss()
                ),
                MDRaisedButton(
                    text="Удалить",
                    md_bg_color=(1, 0, 0, 1),
                    on_release=lambda _: self.perform_delete_records(selected_ids, dialog)
                )
            ]
        )
        dialog.open()

    def perform_delete_records(self, record_ids, dialog):
        """
        Выполняет удаление записей из базы данных

        Args:
            record_ids: Список ID записей для удаления
            dialog: Диалоговое окно подтверждения
        """
        try:
            conn = get_connection()
            cursor = conn.cursor()

            # Удаляем каждую выбранную запись
            for record_id in record_ids:
                delete_record(conn, record_id)

            conn.commit()
            dialog.dismiss()

            # Обновляем список записей
            self.load_story(search_query=self.search_query)
            self.show_message("Успех", f"Удалено записей: {len(record_ids)}")

        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при удалении записей: {str(e)}")
        finally:
            if 'conn' in locals() and conn:
                conn.close()

    def export_to_word(self):
        """
        Экспортирует выбранные записи в документ Word

        Создает документ с подробной информацией и статистикой
        """
        # Проверяем, не гость ли пользователь
        app = MDApp.get_running_app()
        if getattr(app, 'is_guest', False):
            self.show_message("Недоступно", "Экспорт данных недоступен в гостевом режиме")
            return

        # Получаем выбранные записи
        selected_records = self.get_selected_records()

        # Проверяем, есть ли выбранные записи
        if not selected_records:
            self.show_message("Внимание", "Сначала выберите записи для экспорта")
            return

        # Проверяем наличие библиотеки python-docx
        if not DOCX_AVAILABLE:
            self.show_message("Ошибка", "Установите библиотеку: pip install python-docx")
            return

        try:
            # Создаем новый документ Word
            doc = Document()

            # Добавляем заголовок и информацию
            doc.add_heading('Медицинская история записей', 0)
            doc.add_paragraph(f'Отчет создан: {datetime.now().strftime("%d-%m-%Y %H:%M")}')
            doc.add_paragraph(f'Количество записей: {len(selected_records)}')
            doc.add_paragraph()

            # Рассчитываем и добавляем статистику
            numeric_data = self.calculate_statistics(selected_records)
            if numeric_data:
                doc.add_heading('Общая статистика', level=1)
                for stat_text in numeric_data:
                    doc.add_paragraph(stat_text)
                doc.add_paragraph()

            # Добавляем детальную историю записей
            doc.add_heading('Детальная история', level=1)
            for i, record in enumerate(selected_records, 1):
                # Форматируем дату записи
                record_date = self.format_display_date(record[7])
                doc.add_heading(f'Запись {i} от {record_date}', level=2)

                # Форматируем показатели записи
                indicators = self.format_record_indicators(record)
                if indicators:
                    doc.add_paragraph('Показатели: ' + ', '.join(indicators))

                # Добавляем заметки, если они есть
                if record[6] and record[6].strip():
                    doc.add_paragraph(f'Заметки: {record[6]}')

                doc.add_paragraph()

            # Сохраняем документ
            export_dir = self.get_export_directory()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f'medical_history_{timestamp}.docx'
            full_path = os.path.join(export_dir, filename)

            doc.save(full_path)
            self.show_message("Успех", f"Данные экспортированы в Word\nФайл: {filename}")

        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при экспорте в Word: {str(e)}")

    def calculate_statistics(self, records):
        """
        Рассчитывает статистику по записям

        Args:
            records: Список записей

        Returns:
            list: Список текстовых строк со статистикой
        """
        stats = []

        # Статистика по весу
        weights = [self.safe_convert_to_float(r[1]) for r in records if r[1] is not None]
        weights = [w for w in weights if w is not None]
        if weights:
            avg_weight = sum(weights) / len(weights)
            stats.append(f'Средний вес: {avg_weight:.1f} кг')

        # Статистика по давлению
        pressures_sys = [self.safe_convert_to_int(r[2]) for r in records if r[2] is not None]
        pressures_sys = [p for p in pressures_sys if p is not None]
        pressures_dia = [self.safe_convert_to_int(r[3]) for r in records if r[3] is not None]
        pressures_dia = [p for p in pressures_dia if p is not None]
        if pressures_sys and pressures_dia:
            avg_sys = sum(pressures_sys) / len(pressures_sys)
            avg_dia = sum(pressures_dia) / len(pressures_dia)
            stats.append(f'Среднее давление: {avg_sys:.0f}/{avg_dia:.0f} мм рт.ст.')

        # Статистика по пульсу
        pulses = [self.safe_convert_to_int(r[4]) for r in records if r[4] is not None]
        pulses = [p for p in pulses if p is not None]
        if pulses:
            avg_pulse = sum(pulses) / len(pulses)
            stats.append(f'Средний пульс: {avg_pulse:.0f} уд/мин')

        # Статистика по температуре
        temperatures = [self.safe_convert_to_float(r[5]) for r in records if r[5] is not None]
        temperatures = [t for t in temperatures if t is not None]
        if temperatures:
            avg_temp = sum(temperatures) / len(temperatures)
            stats.append(f'Средняя температура: {avg_temp:.1f}°C')

        return stats

    def format_record_indicators(self, record):
        """
        Форматирует показатели записи для отображения

        Args:
            record: Данные записи

        Returns:
            list: Список отформатированных показателей
        """
        indicators = []

        # Вес
        if record[1]:
            weight = self.safe_convert_to_float(record[1])
            if weight is not None:
                indicators.append(f'Вес: {weight} кг')

        # Давление
        if record[2] and record[3]:
            pressure_sys = self.safe_convert_to_int(record[2])
            pressure_dia = self.safe_convert_to_int(record[3])
            if pressure_sys is not None and pressure_dia is not None:
                indicators.append(f'Давление: {pressure_sys}/{pressure_dia}')

        # Пульс
        if record[4]:
            pulse = self.safe_convert_to_int(record[4])
            if pulse is not None:
                indicators.append(f'Пульс: {pulse} уд/мин')

        # Температура
        if record[5]:
            temp = self.safe_convert_to_float(record[5])
            if temp is not None:
                indicators.append(f'Температура: {temp}°C')

        return indicators

    def show_excel_export_dialog(self):
        """
        Показывает диалоговое окно для настройки экспорта в Excel

        Позволяет выбрать тип графика и данные для экспорта
        """
        # Проверяем, не гость ли пользователь
        app = MDApp.get_running_app()
        if getattr(app, 'is_guest', False):
            self.show_message("Недоступно", "Экспорт данных недоступен в гостевом режиме")
            return

        # Получаем выбранные записи
        selected_records = self.get_selected_records()

        # Проверяем, есть ли выбранные записи
        if not selected_records:
            self.show_message("Внимание", "Сначала выберите записи для экспорта")
            return

        # Сбрасываем тип графика на значение по умолчанию
        self.selected_chart_type = "line"

        # Определяем размеры для платформы
        if platform == 'android':
            dialog_width = 0.95  # На Android используем почти всю ширину
            font_style = "Body1"  # Меньший шрифт
            item_height = dp(35)  # Меньшая высота элементов
        else:
            dialog_width = 0.8  # На десктопе 80% ширины
            font_style = "H6"  # Стандартный шрифт
            item_height = dp(40)  # Стандартная высота

        # Создаем содержимое диалогового окна
        dialog_content = MDBoxLayout(
            orientation="vertical",
            spacing=dp(5) if platform == 'android' else dp(10),
            padding=dp(10) if platform == 'android' else dp(15),
            adaptive_height=True
        )

        # Заголовок диалога с адаптивным шрифтом
        dialog_content.add_widget(MDLabel(
            text="Настройка экспорта в Excel",
            theme_text_color="Primary",
            font_style=font_style,
            size_hint_y=None,
            height=dp(35) if platform == 'android' else dp(40),
            halign="center"
        ))

        # Контейнер для выбора типа графика
        chart_type_box = MDBoxLayout(
            orientation="horizontal",
            spacing=dp(5) if platform == 'android' else dp(10),
            size_hint_y=None,
            height=item_height
        )

        # Кнопка выбора типа графика
        self.chart_type_button = MDRaisedButton(
            text="Тип графика: Линейный",
            size_hint_x=0.8 if platform == 'android' else 0.7
        )
        self.chart_type_button.bind(on_release=lambda x: self.open_chart_menu(self.chart_type_button))

        chart_type_box.add_widget(MDLabel(
            text="Тип графика:",
            size_hint_x=0.4,
            size_hint_y=None,
            height=item_height,
            font_style="Body2" if platform == 'android' else "Body1"
        ))
        chart_type_box.add_widget(self.chart_type_button)

        dialog_content.add_widget(chart_type_box)
        dialog_content.add_widget(MDLabel(
            text="Выберите данные для графиков:",
            theme_text_color="Primary",
            size_hint_y=None,
            height=dp(25),
            font_style="Body2" if platform == 'android' else "Body1"
        ))

        # Создаем чекбоксы для выбора данных
        self.weight_check = MDCheckbox(active=True, size_hint=(None, None),
                                       size=(dp(35), dp(35)) if platform == 'android' else (dp(40), dp(40)))
        self.pressure_sys_check = MDCheckbox(active=True, size_hint=(None, None),
                                             size=(dp(35), dp(35)) if platform == 'android' else (dp(40), dp(40)))
        self.pressure_dia_check = MDCheckbox(active=True, size_hint=(None, None),
                                             size=(dp(35), dp(35)) if platform == 'android' else (dp(40), dp(40)))
        self.pulse_check = MDCheckbox(active=True, size_hint=(None, None),
                                      size=(dp(35), dp(35)) if platform == 'android' else (dp(40), dp(40)))
        self.temperature_check = MDCheckbox(active=True, size_hint=(None, None),
                                            size=(dp(35), dp(35)) if platform == 'android' else (dp(40), dp(40)))

        # Добавляем чекбоксы с подписями
        for check, label in [
            (self.weight_check, "Вес"),
            (self.pressure_sys_check, "Систолическое давление"),
            (self.pressure_dia_check, "Диастолическое давление"),
            (self.pulse_check, "Пульс"),
            (self.temperature_check, "Температура")
        ]:
            row = MDBoxLayout(
                orientation="horizontal",
                spacing=dp(5) if platform == 'android' else dp(10),
                size_hint_y=None,
                height=item_height
            )
            row.add_widget(check)
            row.add_widget(MDLabel(
                text=label,
                font_style="Body2" if platform == 'android' else "Body1"
            ))
            dialog_content.add_widget(row)

        # Информация о количестве записей
        dialog_content.add_widget(MDLabel(
            text=f"Будет экспортировано записей: {len(selected_records)}",
            theme_text_color="Secondary",
            size_hint_y=None,
            height=dp(25),
            font_style="Body2" if platform == 'android' else "Body1",
            halign="center"
        ))

        # Создаем диалоговое окно с адаптивными размерами
        self.dialog = MDDialog(
            title="Экспорт в Excel с графиками",
            type="custom",
            content_cls=dialog_content,
            buttons=[
                MDRaisedButton(
                    text="Отмена",
                    size_hint=(0.5, None),
                    height=dp(40),
                    on_release=lambda _: self.dialog.dismiss()
                ),
                MDRaisedButton(
                    text="Экспорт",
                    md_bg_color=(0.2, 0.6, 0.2, 1),
                    size_hint=(0.5, None),
                    height=dp(40),
                    on_release=lambda _: self.perform_excel_export(selected_records)
                ),
            ],
            size_hint=(dialog_width, None),
            height=dp(450) if platform == 'android' else dp(500),
            auto_dismiss=False
        )
        self.dialog.open()

    def perform_excel_export(self, selected_records):
        """
        Выполняет экспорт в Excel с графиками

        Args:
            selected_records: Выбранные записи для экспорта
        """
        # Проверяем наличие библиотеки xlsxwriter
        if not XLSXWRITER_AVAILABLE:
            self.show_message("Ошибка", "Установите библиотеку: pip install xlsxwriter")
            return

        try:
            # Создаем директорию для экспорта
            export_dir = self.get_export_directory()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f'medical_charts_{timestamp}.xlsx'
            full_path = os.path.join(export_dir, filename)

            # Создаем новую книгу Excel
            workbook = xlsxwriter.Workbook(full_path)

            # Создаем лист с данными
            data_worksheet = workbook.add_worksheet('Данные')

            # Заголовки столбцов
            headers = ['Дата', 'Вес (кг)', 'Систолическое давление', 'Диастолическое давление',
                       'Пульс', 'Температура', 'Заметки']
            header_format = workbook.add_format({'bold': True, 'bg_color': '#D3D3D3'})

            # Записываем заголовки
            for col, header in enumerate(headers):
                data_worksheet.write(0, col, header, header_format)

            # Подготавливаем данные для записи
            data_rows = []
            for record in selected_records:
                record_date = self.format_display_date(record[7])
                data_rows.append([
                    record_date,
                    self.safe_convert_to_float(record[1]) or '',
                    self.safe_convert_to_int(record[2]) or '',
                    self.safe_convert_to_int(record[3]) or '',
                    self.safe_convert_to_int(record[4]) or '',
                    self.safe_convert_to_float(record[5]) or '',
                    record[6] or ''
                ])

            # Записываем данные
            for row, data in enumerate(data_rows, start=1):
                for col, value in enumerate(data):
                    data_worksheet.write(row, col, value)

            # Создаем графики
            charts_created = self.create_charts(workbook, data_rows)

            # Закрываем книгу
            workbook.close()

            # Закрываем диалог и показываем сообщение
            self.dialog.dismiss()
            message = f"Данные экспортированы в Excel\nТип графиков: {self.get_chart_type_name(self.selected_chart_type)}"
            if not charts_created:
                message += "\nГрафики не созданы - недостаточно данных"
            self.show_message("Успех", message)

        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при экспорте в Excel: {str(e)}")

    def create_charts(self, workbook, data_rows):
        """
        Создает графики в книге Excel

        Args:
            workbook: Книга Excel
            data_rows: Данные для графиков

        Returns:
            bool: True если создан хотя бы один график, иначе False
        """
        charts_created = False

        # Определяем, какие графики нужно создать
        chart_creators = [
            (self.weight_check, [1], 'График веса', 'Динамика веса', 'Вес (кг)'),
            (self.pressure_sys_check, [2], 'График систолического давления', 'Систолическое давление',
             'Давление (мм рт.ст.)'),
            (self.pressure_dia_check, [3], 'График диастолического давления', 'Диастолическое давление',
             'Давление (мм рт.ст.)'),
            (self.pulse_check, [4], 'График пульса', 'Динамика пульса', 'Пульс (уд/мин)'),
            (self.temperature_check, [5], 'График температуры', 'Динамика температуры', 'Температура (°C)')
        ]

        # Создаем графики для выбранных данных
        for check, columns, sheet_name, title, y_axis in chart_creators:
            if check.active:
                if self.create_single_chart(workbook, data_rows, columns, sheet_name, title, y_axis):
                    charts_created = True

        # Создаем комбинированный график давления
        if self.pressure_sys_check.active and self.pressure_dia_check.active:
            if self.create_single_chart(workbook, data_rows, [2, 3], 'График давления', 'Динамика давления',
                                        'Давление (мм рт.ст.)'):
                charts_created = True

        return charts_created

    def create_single_chart(self, workbook, data_rows, data_columns, sheet_name, title, y_axis_title):
        """
        Создает один график в книге Excel

        Args:
            workbook: Книга Excel
            data_rows: Данные для графика
            data_columns: Индексы столбцов с данными
            sheet_name: Имя листа
            title: Заголовок графика
            y_axis_title: Заголовок оси Y

        Returns:
            bool: True если график создан, иначе False
        """
        try:
            # Фильтруем данные для графика
            chart_data = []
            for row in data_rows:
                date_val = row[0]
                values = []
                valid_row = True

                # Проверяем наличие данных во всех необходимых столбцах
                for col in data_columns:
                    if row[col] == '' or row[col] is None:
                        valid_row = False
                        break
                    values.append(row[col])

                # Добавляем строку, если все данные есть
                if valid_row:
                    chart_data.append([date_val] + values)

            # Проверяем, достаточно ли данных для графика
            if len(chart_data) < 2:
                return False

            # Создаем новый лист для графика
            chart_worksheet = workbook.add_worksheet(sheet_name)
            headers = ['Дата'] + [f'Данные {i + 1}' for i in range(len(data_columns))]
            header_format = workbook.add_format({'bold': True, 'bg_color': '#D3D3D3'})

            # Записываем заголовки
            for col, header in enumerate(headers):
                chart_worksheet.write(0, col, header, header_format)

            # Записываем данные
            for row, data in enumerate(chart_data, start=1):
                for col, value in enumerate(data):
                    chart_worksheet.write(row, col, value)

            # Создаем график
            chart = workbook.add_chart({'type': self.get_xlsxwriter_chart_type()})
            chart.set_title({'name': title})
            chart.set_y_axis({'name': y_axis_title})
            chart.set_x_axis({'name': 'Дата'})

            # Добавляем серии данных
            for i in range(len(data_columns)):
                chart.add_series({
                    'name': headers[i + 1],
                    'categories': [sheet_name, 1, 0, len(chart_data), 0],
                    'values': [sheet_name, 1, i + 1, len(chart_data), i + 1],
                })

            # Вставляем график на лист
            chart_worksheet.insert_chart('D2', chart)
            return True

        except Exception as e:
            print(f"Ошибка при создании графика {title}: {e}")
            return False

    def get_xlsxwriter_chart_type(self):
        """
        Возвращает тип графика для xlsxwriter

        Returns:
            str: Тип графика для xlsxwriter
        """
        chart_types = {
            "line": 'line',
            "bar": 'column',
            "scatter": 'scatter',
            "combo": 'line'  # Для комбинированного используем линейный
        }
        return chart_types.get(self.selected_chart_type, 'line')

    def open_export_folder(self):
        """
        Открывает папку с экспортированными файлами

        Показывает список файлов и позволяет открыть их
        """
        # Получаем директорию экспорта
        export_dir = self.get_export_directory()

        try:
            # Проверяем существование директории
            if not os.path.exists(export_dir):
                os.makedirs(export_dir)

            # Получаем список файлов
            files = os.listdir(export_dir)
            if not files:
                self.show_message("Информация", "Нет экспортированных файлов")
                return

            # Создаем прокручиваемый список файлов
            file_scroll = ScrollView(size_hint=(1, None), size=(dp(400), dp(300)))
            file_list = MDBoxLayout(orientation="vertical", spacing=dp(5), padding=dp(10), adaptive_height=True)

            # Добавляем файлы в список
            for file in sorted(files, reverse=True):  # Сортировка по убыванию (новые сверху)
                file_path = os.path.join(export_dir, file)
                if os.path.isfile(file_path):
                    # Получаем информацию о файле
                    file_size = os.path.getsize(file_path) / 1024  # Размер в КБ
                    file_time = datetime.fromtimestamp(os.path.getctime(file_path)).strftime("%d-%m-%Y %H:%M")

                    # Создаем элемент списка для файла
                    file_item = TwoLineListItem(
                        text=file,
                        secondary_text=f"Размер: {file_size:.1f} KB, Создан: {file_time}"
                    )
                    # Привязываем обработчик открытия файла
                    file_item.bind(on_release=lambda x, fp=file_path: self.open_exported_file(fp))
                    file_list.add_widget(file_item)

            file_scroll.add_widget(file_list)

            # Создаем содержимое диалога
            dialog_content = MDBoxLayout(orientation="vertical", spacing=dp(10), adaptive_height=True)
            dialog_content.add_widget(MDLabel(
                text=f"Найдено файлов: {len(files)}",
                theme_text_color="Secondary",
                size_hint_y=None,
                height=dp(30)
            ))
            dialog_content.add_widget(file_scroll)

            # Создаем диалоговое окно
            self.dialog = MDDialog(
                title="Файлы экспорта",
                type="custom",
                content_cls=dialog_content,
                buttons=[
                    MDRaisedButton(
                        text="Открыть папку",
                        md_bg_color=(0.2, 0.6, 0.2, 1),
                        on_release=lambda _: self.open_folder_in_explorer(export_dir)
                    ),
                    MDRaisedButton(
                        text="Закрыть",
                        on_release=lambda _: self.dialog.dismiss()
                    )
                ],
                size_hint=(0.8, 0.8)
            )
            self.dialog.open()

        except Exception as e:
            self.show_message("Ошибка", f"Не удалось открыть папку: {str(e)}")

    def open_exported_file(self, file_path):
        """
        Открывает экспортированный файл в соответствующем приложении

        Args:
            file_path (str): Путь к файлу
        """
        try:
            # Проверяем существование файла
            if not os.path.exists(file_path):
                self.show_message("Ошибка", f"Файл не найден: {file_path}")
                return

            # Открываем файл в зависимости от платформы
            if sys.platform == 'win32':  # Windows
                # Открываем файл с помощью ассоциированного приложения
                os.startfile(file_path)
            elif sys.platform == 'darwin':  # macOS
                subprocess.run(['open', file_path])
            elif sys.platform.startswith('linux'):  # Linux
                subprocess.run(['xdg-open', file_path])
            elif platform == 'android':  # Android
                self.open_file_android(file_path)
            else:
                # Пытаемся открыть файл стандартным способом
                try:
                    os.startfile(file_path)
                except AttributeError:
                    subprocess.run(['open', file_path], check=False)

            # Закрываем диалог
            if self.dialog:
                self.dialog.dismiss()

        except Exception as e:
            self.show_message("Ошибка", f"Не удалось открыть файл: {str(e)}")

    def open_folder_in_explorer(self, folder_path):
        """
        Открывает папку в проводнике файлов

        Args:
            folder_path (str): Путь к папке
        """
        try:
            # Проверяем существование папки
            if not os.path.exists(folder_path):
                self.show_message("Ошибка", f"Папка не найдена: {folder_path}")
                # Создаем папку, если она не существует
                os.makedirs(folder_path, exist_ok=True)

            # Открываем папку в зависимости от платформы
            if sys.platform == 'win32':  # Windows
                subprocess.run(['explorer', os.path.realpath(folder_path)])
            elif sys.platform == 'darwin':  # macOS
                subprocess.run(['open', folder_path])
            elif sys.platform.startswith('linux'):  # Linux
                subprocess.run(['xdg-open', folder_path])
            elif platform == 'android':  # Android
                # Для Android можно показать путь
                self.show_message("Информация", f"Папка: {folder_path}")
            else:
                subprocess.run(['open', folder_path], check=False)

            # Закрываем диалог
            if self.dialog:
                self.dialog.dismiss()

        except Exception as e:
            self.show_message("Ошибка", f"Не удалось открыть папку: {folder_path} \nОшибка: {str(e)}")

    def open_file_android(self, file_path):
        """
        Открывает файл на Android

        Args:
            file_path (str): Путь к файлу
        """
        try:
            if platform == 'android':
                from jnius import autoclass
                from android import mActivity

                # Используем Android Intent для открытия файла
                Intent = autoclass('android.content.Intent')
                Uri = autoclass('android.net.Uri')
                File = autoclass('java.io.File')

                # Создаем объект файла и URI
                file_obj = File(file_path)
                uri = Uri.fromFile(file_obj)

                # Создаем Intent для просмотра файла
                intent = Intent(Intent.ACTION_VIEW)
                intent.setDataAndType(uri, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                intent.setFlags(Intent.FLAG_ACTIVITY_NEW_TASK)
                intent.addFlags(Intent.FLAG_GRANT_READ_URI_PERMISSION)

                # Запускаем активность
                mActivity.startActivity(intent)
        except Exception as e:
            # Если не удалось открыть, показываем информацию о файле
            self.show_message("Информация", f"Файл: {os.path.basename(file_path)}\nПуть: {file_path}")
