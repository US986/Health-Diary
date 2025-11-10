from kivy.lang import Builder
from kivy.uix.screenmanager import Screen, ScreenManager
from kivy.metrics import dp
from kivy import platform
from kivymd.app import MDApp
from kivymd.uix.dialog import MDDialog
from kivymd.uix.textfield import MDTextField
from kivymd.uix.button import MDRaisedButton, MDFloatingActionButton
from kivymd.uix.boxlayout import MDBoxLayout
from kivymd.uix.list import TwoLineListItem
from kivymd.uix.label import MDLabel
from kivymd.uix.selectioncontrol import MDCheckbox
import os
import subprocess
from datetime import datetime
from base import get_connection
from profile import ProfileScreen
from rules import (
    validate_weight,
    validate_pressure_systolic,
    validate_pressure_diastolic,
    validate_pulse,
    validate_temperature,
    validate_notes,
)

KV = '''
<StoryWindow>:
    name: 'story'
    BoxLayout:
        orientation: "vertical"
        padding: "10dp"
        spacing: "10dp"

        MDLabel:
            text: "История записей"
            halign: "center"
            theme_text_color: "Primary"
            font_style: "H5"
            size_hint_y: None
            height: dp(40)

        ScrollView:
            MDList:
                id: container

        BoxLayout:
            size_hint_y: None
            height: "60dp"
            orientation: "horizontal"
            padding: "8dp"
            spacing: "12dp"
            pos_hint: {"center_x": 0.5}

            MDFloatingActionButton:
                icon: "plus"
                md_bg_color: 0, 0.7, 0, 1
                on_release: root.add_new_record()

            MDFloatingActionButton:
                icon: "account"
                md_bg_color: 0, 0, 1, 1
                on_release: root.on_arrow_pressed()

            MDFloatingActionButton:
                icon: "file-word"
                md_bg_color: 0.2, 0.4, 0.7, 1
                on_release: root.export_to_word()

            MDFloatingActionButton:
                icon: "chart-line"
                md_bg_color: 0.3, 0.6, 0.3, 1
                on_release: root.export_to_excel_with_charts()

            MDFloatingActionButton:
                icon: "folder-open"
                md_bg_color: 0.7, 0.5, 0.2, 1
                on_release: root.open_export_folder()
'''

Builder.load_string(KV)

class DataInputWindow(Screen):
    pass

class StoryWindow(Screen):
    dialog = None
    selected_records = {}

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        if platform == 'android':
            self.request_storage_permission()

    def request_storage_permission(self):
        try:
            if platform == 'android':
                from android.permissions import request_permissions, Permission
                request_permissions([Permission.WRITE_EXTERNAL_STORAGE, Permission.READ_EXTERNAL_STORAGE])
        except:
            pass

    def get_export_directory(self):
        try:
            if platform == 'android':
                from jnius import autoclass
                Environment = autoclass('android.os.Environment')
                downloads_dir = Environment.getExternalStoragePublicDirectory(Environment.DIRECTORY_DOWNLOADS)
                export_dir = os.path.join(downloads_dir.getAbsolutePath(), "HealthDiary")
            else:
                export_dir = os.path.join(os.path.expanduser("~"), "HealthDiary")

            if not os.path.exists(export_dir):
                os.makedirs(export_dir)
            return export_dir
        except Exception as e:
            app_dir = os.path.join(MDApp.get_running_app().user_data_dir, "HealthDiaryExports")
            if not os.path.exists(app_dir):
                os.makedirs(app_dir)
            return app_dir

    def on_pre_enter(self):
        self.load_story()

    def load_story(self):
        user_id = MDApp.get_running_app().get_user_id()
        if not user_id:
            return

        try:
            conn = get_connection()
            with conn.cursor() as cursor:
                cursor.execute("""SELECT id, weight, pressure_systolic, pressure_diastolic,
                                  pulse, temperature, notes, record_date
                                  FROM records WHERE user_id = %s ORDER BY record_date DESC""", (user_id,))
                records = cursor.fetchall()

            story_list = self.ids.container
            story_list.clear_widgets()
            self.selected_records = {}

            if records:
                for record in records:
                    record_id = record[0]
                    record_date = record[7]
                    if isinstance(record_date, str):
                        try:
                            record_date = datetime.strptime(record_date, "%Y-%m-%d").strftime("%d-%m-%Y")
                        except:
                            record_date = record[7]
                    else:
                        record_date = record_date.strftime("%d-%m-%Y")

                    primary_text = f"Дата: {record_date}"
                    secondary_text = (f"Вес: {record[1]} кг, Давление: {record[2]}/{record[3]}, "
                                      f"Пульс: {record[4]}, Темп.: {record[5]}°C")

                    record_container = MDBoxLayout(
                        orientation='horizontal',
                        adaptive_height=True,
                        spacing=dp(10),
                        padding=dp(5)
                    )

                    checkbox = MDCheckbox(
                        size_hint=(None, None),
                        size=(dp(40), dp(40)),
                        active=False
                    )

                    list_item = TwoLineListItem(
                        text=primary_text,
                        secondary_text=secondary_text
                    )

                    list_item.bind(on_release=lambda x, rec=record: self.open_edit_form(rec))

                    checkbox.bind(
                        active=lambda instance, value, rec_id=record_id: self.on_checkbox_active(instance, value,
                                                                                                 rec_id))

                    record_container.add_widget(checkbox)
                    record_container.add_widget(list_item)

                    story_list.add_widget(record_container)

                    self.selected_records[record_id] = {
                        'container': record_container,
                        'checkbox': checkbox,
                        'list_item': list_item,
                        'record': record,
                        'selected': False
                    }
            else:
                empty_label = MDLabel(
                    text="История записей пуста",
                    halign="center",
                    theme_text_color="Secondary",
                    font_style="H6",
                    size_hint_y=None,
                    height=dp(60)
                )
                story_list.add_widget(empty_label)

        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при загрузке истории: {str(e)}")
        finally:
            if 'conn' in locals() and conn:
                conn.close()

    def on_checkbox_active(self, checkbox, value, record_id):
        if record_id in self.selected_records:
            record_data = self.selected_records[record_id]
            record_data['selected'] = value

            if value:
                record_data['list_item'].md_bg_color = (0.9, 0.9, 1, 0.3)
            else:
                record_data['list_item'].md_bg_color = (1, 1, 1, 1)

    def get_selected_records(self):
        selected = []
        for record_id, record_data in self.selected_records.items():
            if record_data['selected']:
                selected.append(record_data['record'])
        return selected

    def add_new_record(self):
        self.manager.current = "options"

    def open_edit_form(self, record):
        self.edit_record_dialog(record)

    def edit_record_dialog(self, record):
        dialog_box = MDBoxLayout(
            orientation="vertical",
            spacing=dp(10),
            padding=dp(10),
            adaptive_height=True
        )

        weight_input = MDTextField(
            hint_text="Вес",
            text=str(record[1]) if record[1] is not None else "",
            mode="rectangle"
        )
        pressure_systolic_input = MDTextField(
            hint_text="Систолическое давление",
            text=str(record[2]) if record[2] is not None else "",
            mode="rectangle"
        )
        pressure_diastolic_input = MDTextField(
            hint_text="Диастолическое давление",
            text=str(record[3]) if record[3] is not None else "",
            mode="rectangle"
        )
        pulse_input = MDTextField(
            hint_text="Пульс",
            text=str(record[4]) if record[4] is not None else "",
            mode="rectangle"
        )
        temperature_input = MDTextField(
            hint_text="Температура",
            text=str(record[5]) if record[5] is not None else "",
            mode="rectangle"
        )
        notes_input = MDTextField(
            hint_text="Заметки",
            text=str(record[6]) if record[6] is not None else "",
            mode="rectangle",
            multiline=True
        )

        dialog_box.add_widget(weight_input)
        dialog_box.add_widget(pressure_systolic_input)
        dialog_box.add_widget(pressure_diastolic_input)
        dialog_box.add_widget(pulse_input)
        dialog_box.add_widget(temperature_input)
        dialog_box.add_widget(notes_input)

        self.dialog = MDDialog(
            title="Редактировать запись",
            type="custom",
            content_cls=dialog_box,
            buttons=[
                MDRaisedButton(
                    text="Отмена",
                    md_bg_color=(0.7, 0.7, 0.7, 1),
                    on_release=lambda _: self.dialog.dismiss()
                ),
                MDRaisedButton(
                    text="Сохранить",
                    md_bg_color=(0, 0, 1, 1),
                    on_release=lambda _: self.save_edit_record(
                        record[0], weight_input.text, pressure_systolic_input.text,
                        pressure_diastolic_input.text, pulse_input.text,
                        temperature_input.text, notes_input.text
                    )
                ),
            ],
        )
        self.dialog.open()

    def save_edit_record(self, record_id, weight, pressure_systolic, pressure_diastolic, pulse, temperature, notes):
        try:
            weight = validate_weight(weight)
            pressure_systolic = validate_pressure_systolic(pressure_systolic)
            pressure_diastolic = validate_pressure_diastolic(pressure_diastolic)
            pulse = validate_pulse(pulse)
            temperature = validate_temperature(temperature)
            notes = validate_notes(notes)

            conn = get_connection()
            with conn.cursor() as cursor:
                cursor.execute("""
                    UPDATE records
                    SET weight=%s, pressure_systolic=%s,
                        pressure_diastolic=%s, pulse=%s, temperature=%s, notes=%s
                    WHERE id=%s
                """, (weight, pressure_systolic, pressure_diastolic, pulse, temperature, notes, record_id))
                conn.commit()

            self.dialog.dismiss()
            self.load_story()
            self.show_message("Успех", "Запись успешно обновлена")

        except ValueError as ve:
            self.show_message("Ошибка ввода", str(ve))
        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при сохранении данных: {str(e)}")
        finally:
            if 'conn' in locals() and conn:
                conn.close()

    def show_message(self, title, text):
        dialog = MDDialog(
            title=title,
            text=text,
            buttons=[
                MDRaisedButton(
                    text="OK",
                    md_bg_color=(0, 0, 1, 1),
                    on_release=lambda _: dialog.dismiss()
                )
            ]
        )
        dialog.open()

    def on_arrow_pressed(self):
        self.manager.current = "profile"

    def format_record_date(self, record_date):
        if isinstance(record_date, str):
            try:
                if '-' in record_date:
                    parts = record_date.split('-')
                    if len(parts) == 3:
                        if len(parts[0]) == 4:
                            return f"{parts[2]}-{parts[1]}-{parts[0]}"
                        else:
                            return record_date
            except:
                return record_date
        else:
            return record_date.strftime("%d-%m-%Y")
        return record_date

    def export_to_word(self):
        selected_records = self.get_selected_records()

        if not selected_records:
            self.show_message("Внимание", "Сначала выберите записи для экспорта")
            return

        try:
            from docx import Document

            doc = Document()

            doc.add_heading('Медицинская история записей', 0)
            doc.add_paragraph(f'Отчет создан: {datetime.now().strftime("%d-%m-%Y %H:%M")}')
            doc.add_paragraph(f'Количество выбранных записей: {len(selected_records)}')
            doc.add_paragraph()

            doc.add_heading('Общая информация', level=1)
            doc.add_paragraph()
            doc.add_heading('Анализ показателей', level=1)

            weights = [r[1] for r in selected_records if r[1] is not None]
            if weights:
                avg_weight = sum(weights) / len(weights)
                doc.add_paragraph(f'Средний вес: {avg_weight:.1f} кг')

            pressures_sys = [r[2] for r in selected_records if r[2] is not None]
            pressures_dia = [r[3] for r in selected_records if r[3] is not None]
            if pressures_sys and pressures_dia:
                avg_sys = sum(pressures_sys) / len(pressures_sys)
                avg_dia = sum(pressures_dia) / len(pressures_dia)
                doc.add_paragraph(f'Среднее давление: {avg_sys:.0f}/{avg_dia:.0f} мм рт.ст.')

            doc.add_heading('Детальная история записей', level=1)

            for i, record in enumerate(selected_records, 1):
                record_date = self.format_record_date(record[7])
                beautiful_date = self.format_date_beautiful(record_date)

                doc.add_heading(f'Запись {i} от {beautiful_date}', level=2)

                indicators = []
                if record[1]:
                    indicators.append(f'Вес: {record[1]} кг')
                if record[2] and record[3]:
                    indicators.append(f'Давление: {record[2]}/{record[3]}')
                if record[4]:
                    indicators.append(f'Пульс: {record[4]}')
                if record[5]:
                    indicators.append(f'Температура: {record[5]}°C')

                if indicators:
                    doc.add_paragraph('Основные показатели: ' + ', '.join(indicators))

                if record[6] and record[6].strip():
                    notes_para = doc.add_paragraph()
                    notes_para.add_run('Заметки: ').bold = True
                    notes_para.add_run(record[6])

                doc.add_paragraph()

            export_dir = self.get_export_directory()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f'medical_history_{timestamp}.docx'
            full_path = os.path.join(export_dir, filename)

            doc.save(full_path)
            self.show_message("Успех", f"Данные экспортированы в Word\nФайл: {filename}")

        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при экспорте в Word: {str(e)}")

    def export_to_excel_with_charts(self):
        selected_records = self.get_selected_records()

        if not selected_records:
            self.show_message("Внимание", "Сначала выберите записи для экспорта")
            return

        try:
            import openpyxl
            from openpyxl import Workbook

            wb = Workbook()
            ws_data = wb.active
            ws_data.title = "Данные"

            headers = ['Дата', 'Вес (кг)', 'Систолическое давление', 'Диастолическое давление',
                       'Пульс', 'Температура', 'Заметки']
            ws_data.append(headers)

            for record in selected_records:
                record_date = self.format_record_date(record[7])
                ws_data.append([
                    record_date,
                    record[1] if record[1] is not None else '',
                    record[2] if record[2] is not None else '',
                    record[3] if record[3] is not None else '',
                    record[4] if record[4] is not None else '',
                    record[5] if record[5] is not None else '',
                    record[6] if record[6] is not None else ''
                ])

            try:
                from openpyxl.chart import LineChart, Reference

                ws_charts = wb.create_sheet("Графики")

                weight_data = [r[1] for r in selected_records if r[1] is not None]
                if len(weight_data) >= 2:
                    chart1 = LineChart()
                    chart1.title = "Динамика веса"
                    chart1.style = 13
                    chart1.y_axis.title = 'Вес (кг)'

                    data = Reference(ws_data, min_col=2, min_row=2, max_row=len(selected_records) + 1)
                    chart1.add_data(data, titles_from_data=True)

                    dates = Reference(ws_data, min_col=1, min_row=2, max_row=len(selected_records) + 1)
                    chart1.set_categories(dates)

                    ws_charts.add_chart(chart1, "A1")
            except Exception:
                pass

            export_dir = self.get_export_directory()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f'medical_charts_{timestamp}.xlsx'
            full_path = os.path.join(export_dir, filename)

            wb.save(full_path)
            self.show_message("Успех", "Данные экспортированы в Excel")

        except Exception as e:
            self.show_message("Ошибка", f"Ошибка при экспорте в Excel: {str(e)}")

    def format_date_beautiful(self, date_str):
        try:
            months = {
                '01': 'января', '02': 'февраля', '03': 'марта',
                '04': 'апреля', '05': 'мая', '06': 'июня',
                '07': 'июля', '08': 'августа', '09': 'сентября',
                '10': 'октября', '11': 'ноября', '12': 'декабря'
            }

            parts = date_str.split('-')
            if len(parts) == 3:
                day, month, year = parts
                if month in months:
                    return f"{day} {months[month]} {year} года"

            return date_str
        except:
            return date_str

    def open_export_folder(self):
        export_dir = self.get_export_directory()

        try:
            files = os.listdir(export_dir)
            if not files:
                self.show_message("Информация", "Нет экспортированных файлов")
                return

            file_list = MDBoxLayout(
                orientation="vertical",
                spacing=dp(10),
                padding=dp(10),
                adaptive_height=True
            )

            for file in sorted(files, reverse=True):
                file_path = os.path.join(export_dir, file)
                file_size = os.path.getsize(file_path) / 1024
                file_time = datetime.fromtimestamp(os.path.getctime(file_path)).strftime("%d-%m-%Y %H:%M")

                file_item = TwoLineListItem(
                    text=file,
                    secondary_text=f"Размер: {file_size:.1f} KB, Создан: {file_time}"
                )

                file_item.bind(on_release=lambda x, fp=file_path, fname=file: self.open_exported_file(fp, fname))
                file_list.add_widget(file_item)

            self.dialog = MDDialog(
                title="Файлы экспорта",
                type="custom",
                content_cls=file_list,
                buttons=[
                    MDRaisedButton(
                        text="Открыть папку в проводнике",
                        md_bg_color=(0.2, 0.6, 0.2, 1),
                        on_release=lambda _: self.open_folder_in_explorer(export_dir)
                    ),
                    MDRaisedButton(
                        text="Закрыть",
                        md_bg_color=(0.7, 0.7, 0.7, 1),
                        on_release=lambda _: self.dialog.dismiss()
                    )
                ],
            )
            self.dialog.open()

        except Exception as e:
            self.show_message("Ошибка", "Не удалось открыть папку")

    def open_exported_file(self, file_path, file_name):
        try:
            if platform == 'win':
                os.startfile(file_path)
            elif platform == 'linux':
                subprocess.run(['xdg-open', file_path])
            elif platform == 'macosx':
                subprocess.run(['open', file_path])
            elif platform == 'android':
                self.open_file_android(file_path, file_name)
            else:
                subprocess.run(['open', file_path], check=False)

            self.dialog.dismiss()

        except Exception as e:
            self.show_message("Ошибка", f"Не удалось открыть файл: {str(e)}")

    def open_folder_in_explorer(self, folder_path):
        try:
            if platform == 'win':
                os.startfile(folder_path)
            elif platform == 'linux':
                subprocess.run(['xdg-open', folder_path])
            elif platform == 'macosx':
                subprocess.run(['open', folder_path])
            elif platform == 'android':
                self.show_message("Информация", f"Папка экспорта: {folder_path}")
            else:
                subprocess.run(['open', folder_path], check=False)

            self.dialog.dismiss()

        except Exception as e:
            self.show_message("Ошибка", f"Не удалось открыть папку: {str(e)}")

    def open_file_android(self, file_path, file_name):
        try:
            from jnius import autoclass
            from android import mActivity

            Intent = autoclass('android.content.Intent')
            Uri = autoclass('android.net.Uri')
            File = autoclass('java.io.File')

            intent = Intent()
            intent.setAction(Intent.ACTION_VIEW)

            if file_name.lower().endswith('.docx'):
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif file_name.lower().endswith('.xlsx'):
                mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            else:
                mime_type = "*/*"

            file_obj = File(file_path)
            uri = Uri.fromFile(file_obj)

            intent.setDataAndType(uri, mime_type)
            intent.setFlags(Intent.FLAG_ACTIVITY_NEW_TASK)
            intent.addFlags(Intent.FLAG_GRANT_READ_URI_PERMISSION)

            mActivity.startActivity(intent)

        except Exception as e:
            self.show_message("Информация", f"Файл: {file_name}\nПуть: {file_path}")

class HealthDiaryApp(MDApp):
    user_id = None

    def build(self):
        sm = ScreenManager()
        sm.add_widget(DataInputWindow(name="data_input"))
        sm.add_widget(StoryWindow(name="story"))
        sm.add_widget(ProfileScreen(name="profile"))
        sm.current = "data_input"
        return sm

    def get_user_id(self):
        return self.user_id

    def set_user_id(self, user_id):
        self.user_id = user_id

    def go_to_story(self):
        self.root.current = "story"

if __name__ == "__main__":
    HealthDiaryApp().run()





















